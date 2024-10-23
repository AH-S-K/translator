import os
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from PyQt5.QtWidgets import QApplication, QFileDialog
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import fitz
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from tqdm import tqdm
import ctypes
def extract_text_from_pdf(pdf_file_path):
    if not os.path.exists(pdf_file_path):
        print(f"Error: The file {pdf_file_path} does not exist.")
        return None
    try:
        document = fitz.open(pdf_file_path)
        total_pages = document.page_count
        pages_text = []
        for page_number in range(total_pages):
            page = document.load_page(page_number)
            text = page.get_text()
            text = re.sub(r'Copyright.*\n?', '', text)
            text = re.sub(r'(SUT|IE|Information Technology|Turban Industry 4\.0)[\s\-]*\n?', '', text)
            text = re.sub(r'\n(?![.!؟!?\d\)\*\-])', ' ', text)
            bullet_patterns = r'(\n?(\u2022|\u2013|\u25cb|\u25a0|\u2666|\u2023|\u203a|\u261b|\u2192|\u25cf|\u25e6|\u25aa|\u25ab|\u25c9|\u25d8|\u25d9|\u25c6|\u25c7|\u274f|\u2756|\u2765|\u2767))'
            text = re.sub(bullet_patterns, r'\n\2', text)
            text = re.sub(r'([A-Z])\s+([A-Z])', r'\1\2', text)
            text = re.sub(r' +', ' ', text)
            pages_text.append(text)
        print("text extracted successfuly")
        return pages_text
    except Exception as e:
        print(e)
        return None
    finally:
        if document is not None:
            document.close()
def translate_text(pages_text):
    options = Options()
    options.add_argument("--incognito")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    user32 = ctypes.windll.user32
    screen_width = user32.GetSystemMetrics(0)
    screen_height = user32.GetSystemMetrics(1)
    chrome_window_x = int(screen_width * 0.4)
    chrome_window_y = int(screen_height * 0.01)
    options.add_argument(f"window-position={chrome_window_x},{chrome_window_y}")
    options.add_argument(f"window-size={int(screen_width * 0.4)},{int(screen_height * 0.8)}")
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=options)
    driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
        "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
    })
    translations = []
    try:
        driver.get("https://translate.google.com/?sl=en&tl=fa&op=translate")
        for page in tqdm(pages_text, desc="Translating pages", unit="page"):
            if not page.strip():
                translations.append("")
                continue
            input_box = WebDriverWait(driver, 15).until(
                EC.presence_of_element_located((By.XPATH, "//textarea[@aria-label='Source text']"))
            )
            input_box.clear()
            input_box.send_keys(page)
            input_box.send_keys(Keys.RETURN)
            end_time_clear = time.time() + 5
            while time.time() < end_time_clear:
                translated_elements = driver.find_elements(By.XPATH, "//span[@jsname='W297wb']")
                if not translated_elements:
                    break
                time.sleep(0.25)
            translated_text = ""
            end_time = time.time() + 15
            stable_translation_time = 5
            last_text = ""
            stable_start = time.time()
            while time.time() < end_time:
                translated_elements = driver.find_elements(By.XPATH, "//span[@jsname='W297wb']")
                translated_text = '\n'.join([element.text for element in translated_elements])
                if translated_text == last_text:
                    if time.time() - stable_start >= stable_translation_time:
                        break
                else:
                    last_text = translated_text
                    stable_start = time.time()
                time.sleep(0.5)
            translations.append(translated_text)
        print("Translation successful!")
        return translations
    except Exception as e:
        print(e)
        return translations
    finally:
        driver.quit()
def pdf_to_word_with_translations(pdf_path, translations, output_docx_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        image_path = f"temp_page_{page_number + 1}.png"
        pix.save(image_path)
        doc.add_picture(image_path, width=Inches(6))
        if translations and page_number < len(translations):
            paragraph = doc.add_paragraph(translations[page_number])
        doc.add_page_break()
        os.remove(image_path)
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for run in paragraph.runs:
            run_element = run._element
            rtl_element = OxmlElement('w:rtl')
            rtl_element.set(qn('w:val'), '1')
            run_element.get_or_add_rPr().append(rtl_element)
    doc.save(output_docx_path)
if __name__ == "__main__":
    app = QApplication([])
    pdf_file_path, _ = QFileDialog.getOpenFileName(None, "Select PDF File", "", "PDF Files (*.pdf)")
    app.quit()
    if pdf_file_path:
        pages_text = extract_text_from_pdf(pdf_file_path)
        if pages_text:
            translated_pages = translate_text(pages_text)
            if translated_pages:
                desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                output_docx_name = os.path.splitext(os.path.basename(pdf_file_path))[0] + "_translated.docx"
                output_docx_path = os.path.join(desktop_path, output_docx_name)
                pdf_to_word_with_translations(pdf_file_path, translated_pages, output_docx_path)
                print(f"Word document created successfully: {output_docx_path}")
